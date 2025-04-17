# import os
# from io import BytesIO
# from langchain.chains import LLMChain
# from langchain.prompts import PromptTemplate
# from langchain_groq import ChatGroq
# from langchain.text_splitter import RecursiveCharacterTextSplitter
# from docx import Document
# import subprocess
# from pathlib import Path
# import shutil
# import traceback

# def initialize_groq_llm():
#     api_key = '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'
#     if not api_key:
#         raise ValueError("GROQ_API_KEY is not set in the environment variables.")
#     return ChatGroq(api_key=api_key, model="llama-3.3-70b-versatile")  # Updated to use the specified model

# # Update the analyze_srs_document function to handle .docx files
# def analyze_srs_document(srs_document):
#     """
#     Analyze the SRS document (text or image) to extract key software requirements.

#     Args:
#         srs_document: The input SRS document (binary content of a .docx file).

#     Returns:
#         dict: Extracted requirements including API endpoints, database schema, etc.
#     """
#     # Read the .docx file content
#     try:
#         # Wrap the binary content in a BytesIO stream
#         doc = Document(BytesIO(srs_document))
#         # Clean up the extracted text by removing extra whitespace and tabs
#         content = "\n".join([paragraph.text.strip().replace("\t", " ") for paragraph in doc.paragraphs if paragraph.text.strip()])
#     except Exception as e:
#         raise ValueError(f"Error reading .docx file: {str(e)}")

#     # Split the document into smaller chunks
#     text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=200)
#     chunks = text_splitter.split_text(content)

#     # Initialize the Groq LLM
#     llm = initialize_groq_llm()

#     # Process each chunk and collect results for each section separately
#     formatted_results = {
#         "api_endpoints": [],
#         "backend_logic": [],
#         "database_schema": [],
#         "authentication_requirements": []
#     }

#     # Refine the section-specific prompts to instruct the LLM to write from a requirements perspective
#     section_prompts = {
#         "api_endpoints": "Extract all API endpoints and their parameters from the document. Write the endpoints as requirements, specifying what each endpoint should do, the HTTP method, and the required parameters.",
#         "backend_logic": "Extract all backend logic (business rules, required computations) from the document. Write the logic as requirements, specifying what the system must do and the rules it must enforce.",
#         "database_schema": "Extract the database schema (tables, relationships, constraints) from the document. Write the schema as requirements, specifying what tables must exist, their columns, data types, and relationships.",
#         "authentication_requirements": "Extract all authentication and authorization requirements from the document. Write the requirements specifying what authentication methods, roles, and permissions the system must implement."
#     }

#     for section, section_prompt in section_prompts.items():
#         section_results = []
#         for chunk in chunks:
#             # Update the prompt for the specific section
#             prompt = PromptTemplate(
#                 input_variables=["document"],
#                 template=f"""
#                 You are an expert software engineer. Analyze the following Software Requirements Specification (SRS) document and {section_prompt}

#                 Document:
#                 {{document}}
#                 """
#             )

#             # Create a LangChain LLMChain
#             chain = LLMChain(llm=llm, prompt=prompt)

#             # Run the chain for the current chunk
#             result = chain.run(document=chunk)
#             section_results.append(result)

#         # Combine results for the section
#         formatted_results[section] = [
#             result.replace("\n", " ").replace("\t", " ").strip() for result in section_results if result.strip()
#         ]

#     # Ensure all sections are populated
#     for key in formatted_results:
#         if not formatted_results[key]:
#             formatted_results[key] = ["No data extracted"]

#     return formatted_results

# # Define a custom task-based execution pipeline to automate project generation
# def execute_pipeline(requirements):
#     """
#     Custom pipeline to automate project generation based on extracted requirements.

#     Args:
#         requirements (dict): Extracted requirements including API endpoints, backend logic, database schema, and authentication requirements.

#     Returns:
#         None
#     """
#     try:
#         print("Step 1: Analyzing SRS Document...")
#         extracted_requirements = analyze_srs_document(requirements.get("srs_document"))
#         print("SRS Document analyzed successfully.")

#         print("Step 2: Generating Code...")
#         generate_code(extracted_requirements)
#         print("Code generation completed successfully.")

#         print("Step 3: Setting Up Environment...")
#         setup_environment()
#         print("Environment setup completed successfully.")

#         print("Step 4: Setting Up Database...")
#         setup_database()
#         print("Database setup completed successfully.")

#         print("Step 5: Validating Project...")
#         validate_project()
#         print("Project validation completed successfully.")

#     except Exception as e:
#         print(f"Error during pipeline execution: {str(e)}")
#         raise

# # Placeholder functions for each node
# def generate_code(requirements):
#     print("Generating code based on requirements...")
#     project_root = Path("c:/study/autogen/generated_project")
    
#     # Check if the project directory exists and remove it if it does
#     project_dir = Path("c:/study/autogen/generated_project")
#     if project_dir.exists():
#         print("Existing project directory found. Removing it...")
#         shutil.rmtree(project_dir)
#         print("Existing project directory removed.")

#     project_root.mkdir(exist_ok=True)

#     # Define the folder structure
#     folders = [
#         "app/api/routes",
#         "app/models",
#         "app/services",
#         "app/core",
#         "tests"
#     ]

#     for folder in folders:
#         (project_root / folder).mkdir(parents=True, exist_ok=True)

#     # Initialize the Groq LLM
#     llm = initialize_groq_llm()

#     # Ensure the input to `llm.generate` is properly formatted as a list of dictionaries
#     main_prompt = [{"role": "system", "content": "Generate the main entry point for a FastAPI application."}]
#     main_code = llm.generate(messages=main_prompt)
#     if main_code.strip():
#         (project_root / "app/main.py").write_text(main_code)
#     else:
#         print("Error: LLM did not generate code for main.py")

#     # Similarly update other prompts to be lists of dictionaries
#     routes_prompt = [{"role": "system", "content": "Generate API routes based on the extracted requirements."}]
#     routes_code = llm.generate(messages=routes_prompt)
#     if routes_code.strip():
#         (project_root / "app/api/routes/__init__.py").write_text(routes_code)
#     else:
#         print("Error: LLM did not generate code for routes.")

#     models_prompt = [{"role": "system", "content": "Generate Pydantic models based on the extracted requirements."}]
#     models_code = llm.generate(messages=models_prompt)
#     if models_code.strip():
#         (project_root / "app/models/__init__.py").write_text(models_code)
#     else:
#         print("Error: LLM did not generate code for models.")

#     services_prompt = [{"role": "system", "content": "Generate service functions based on the extracted requirements."}]
#     services_code = llm.generate(messages=services_prompt)
#     if services_code.strip():
#         (project_root / "app/services/__init__.py").write_text(services_code)
#     else:
#         print("Error: LLM did not generate code for services.")

#     config_prompt = [{"role": "system", "content": "Generate a configuration file for the FastAPI project."}]
#     config_code = llm.generate(messages=config_prompt)
#     if config_code.strip():
#         (project_root / "app/core/config.py").write_text(config_code)
#     else:
#         print("Error: LLM did not generate code for config.")

#     print("Code generation completed successfully.")


# def setup_environment():
#     print("Setting up virtual environment and installing dependencies...")
#     project_root = Path("c:/study/autogen/generated_project")

#     # Create a virtual environment
#     venv_path = project_root / "venv"
#     try:
#         print(f"Creating virtual environment at {venv_path}...")
#         process = subprocess.Popen(
#             ["python", "-m", "venv", str(venv_path)],
#             stdout=subprocess.PIPE,
#             stderr=subprocess.PIPE
#         )
#         stdout, stderr = process.communicate()
#         if process.returncode != 0:
#             print(f"Error creating virtual environment: {stderr.decode()}")
#             raise Exception(stderr.decode())
#         print(stdout.decode())
#     except Exception as e:
#         print(f"Error creating virtual environment: {e}")
#         print(traceback.format_exc())
#         raise

#     # Activate the virtual environment and install dependencies
#     pip_path = venv_path / "Scripts" / "pip"
#     requirements_file = Path("c:/study/autogen/srs-analyzer/requirements.txt")

#     try:
#         if requirements_file.exists():
#             print("Installing dependencies from requirements.txt...")
#             process = subprocess.Popen(
#                 [str(pip_path), "install", "-r", str(requirements_file)],
#                 stdout=subprocess.PIPE,
#                 stderr=subprocess.PIPE
#             )
#             stdout, stderr = process.communicate()
#             if process.returncode != 0:
#                 print(f"Error installing dependencies: {stderr.decode()}")
#                 raise Exception(stderr.decode())
#             print(stdout.decode())
#         else:
#             print("requirements.txt not found. Installing default dependencies...")
#             default_dependencies = ["fastapi", "uvicorn", "psycopg2", "alembic"]
#             process = subprocess.Popen(
#                 [str(pip_path), "install"] + default_dependencies,
#                 stdout=subprocess.PIPE,
#                 stderr=subprocess.PIPE
#             )
#             stdout, stderr = process.communicate()
#             if process.returncode != 0:
#                 print(f"Error installing default dependencies: {stderr.decode()}")
#                 raise Exception(stderr.decode())
#             print(stdout.decode())
#     except Exception as e:
#         print(f"Error installing dependencies: {e}")
#         print(traceback.format_exc())
#         raise

#     print("Environment setup and dependency installation completed successfully.")


# def setup_database():
#     print("Setting up PostgreSQL database and migrations...")
#     project_root = Path("c:/study/autogen/generated_project")

#     # Create Alembic configuration
#     alembic_ini = project_root / "alembic.ini"
#     alembic_ini.write_text(
#         """[alembic]
# script_location = migrations
# sqlalchemy.url = postgresql+psycopg://vectordb:vectordb@localhost:5432/vectordb

# [loggers]
# keys = root,sqlalchemy,alembic

# [handlers]
# keys = console

# [formatters]
# keys = generic

# [logger_root]
# level = WARN
# handlers = console
# qualname =

# [logger_sqlalchemy]
# level = WARN
# handlers = console
# qualname = sqlalchemy.engine

# [logger_alembic]
# level = INFO
# handlers = console
# qualname = alembic

# [handler_console]
# class = StreamHandler
# args = (sys.stderr,)
# level = NOTSET
# formatter = generic

# [formatter_generic]
# format = %(asctime)s %(levelname)-5.5s [%(name)s] %(message)s
# """
#     )

#     # Initialize Alembic
#     migrations_path = project_root / "migrations"
#     subprocess.run(["alembic", "init", str(migrations_path)], check=True, cwd=str(project_root))

#     print("Database setup completed successfully.")


# def validate_project():
#     print("Validating the generated project...")
#     project_root = Path("c:/study/autogen/generated_project")

#     # Run the FastAPI application on a different port to avoid conflicts
#     main_file = project_root / "app/main.py"
#     if main_file.exists():
#         print("Starting the generated project on localhost:8001...")
#         try:
#             process = subprocess.Popen(
#                 ["uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", "8001", "--reload"],
#                 cwd=str(project_root),
#                 stdout=subprocess.PIPE,
#                 stderr=subprocess.PIPE
#             )
#             stdout, stderr = process.communicate()
#             if process.returncode != 0:
#                 print(f"Error starting the project: {stderr.decode()}")
#                 raise Exception(stderr.decode())
#             print(stdout.decode())
#             print("Project is running at http://127.0.0.1:8001")
#         except Exception as e:
#             print(f"Error starting the project: {e}")
#             print(traceback.format_exc())
#             raise
#     else:
#         print("Main file not found. Validation failed.")
#         raise FileNotFoundError("Main file not found in the generated project.")

#     print("Project validation completed successfully.")









# Code Generated by Sidekick is for learning and experimentation purposes only.

import os
from io import BytesIO
from langchain.chains import LLMChain
from langchain.prompts import PromptTemplate
from langchain_groq import ChatGroq
from langchain.text_splitter import RecursiveCharacterTextSplitter
from docx import Document
import subprocess
from pathlib import Path
import shutil
import traceback

def initialize_groq_llm():
    api_key = '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'
    if not api_key:
        raise ValueError("GROQ_API_KEY is not set in the environment variables.")
    return ChatGroq(api_key=api_key, model="llama-3.3-70b-versatile")  # Updated to use the specified model

def analyze_srs_document(srs_document):
    """
    Analyze the SRS document (text or image) to extract key software requirements.

    Args:
        srs_document: The input SRS document (binary content of a .docx file).

    Returns:
        dict: Extracted requirements including API endpoints, database schema, etc.
    """
    try:
        doc = Document(BytesIO(srs_document))
        content = "\n".join([paragraph.text.strip().replace("\t", " ") for paragraph in doc.paragraphs if paragraph.text.strip()])
    except Exception as e:
        raise ValueError(f"Error reading .docx file: {str(e)}")

    text_splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=200)
    chunks = text_splitter.split_text(content)

    llm = initialize_groq_llm()

    formatted_results = {
        "api_endpoints": [],
        "backend_logic": [],
        "database_schema": [],
        "authentication_requirements": []
    }

    section_prompts = {
        "api_endpoints": "Extract all API endpoints and their parameters from the document. Write the endpoints as requirements, specifying what each endpoint should do, the HTTP method, and the required parameters.",
        "backend_logic": "Extract all backend logic (business rules, required computations) from the document. Write the logic as requirements, specifying what the system must do and the rules it must enforce.",
        "database_schema": "Extract the database schema (tables, relationships, constraints) from the document. Write the schema as requirements, specifying what tables must exist, their columns, data types, and relationships.",
        "authentication_requirements": "Extract all authentication and authorization requirements from the document. Write the requirements specifying what authentication methods, roles, and permissions the system must implement."
    }

    for section, section_prompt in section_prompts.items():
        section_results = []
        for chunk in chunks:
            prompt = PromptTemplate(
                input_variables=["document"],
                template=f"""
                You are an expert software engineer. Analyze the following Software Requirements Specification (SRS) document and {section_prompt}

                Document:
                {{document}}
                """
            )

            chain = LLMChain(llm=llm, prompt=prompt)
            result = chain.invoke(document=chunk)
            section_results.append(result)

        formatted_results[section] = [
            result.replace("\n", " ").replace("\t", " ").strip() for result in section_results if result.strip()
        ]

    for key in formatted_results:
        if not formatted_results[key]:
            formatted_results[key] = ["No data extracted"]

    return formatted_results

def execute_pipeline(requirements):
    """
    Custom pipeline to automate project generation based on extracted requirements.

    Args:
        requirements (dict): Extracted requirements including API endpoints, backend logic, database schema, and authentication requirements.

    Returns:
        None
    """
    try:
        print("Step 1: Analyzing SRS Document...")
        extracted_requirements = analyze_srs_document(requirements.get("srs_document"))
        print("SRS Document analyzed successfully.")

        print("Step 2: Generating Code...")
        generate_code(extracted_requirements)
        print("Code generation completed successfully.")

        print("Step 3: Setting Up Environment...")
        setup_environment()
        print("Environment setup completed successfully.")

        print("Step 4: Setting Up Database...")
        setup_database()
        print("Database setup completed successfully.")

        print("Step 5: Validating Project...")
        validate_project()
        print("Project validation completed successfully.")

    except Exception as e:
        print(f"Error during pipeline execution: {str(e)}")
        raise

def generate_code(requirements):
    print("Generating code based on requirements...")
    project_root = Path("c:/study/autogen/generated_project")

    project_dir = Path("c:/study/autogen/generated_project")
    if project_dir.exists():
        print("Existing project directory found. Removing it...")
        shutil.rmtree(project_dir)
        print("Existing project directory removed.")

    project_root.mkdir(exist_ok=True)

    folders = [
        "app/api/routes",
        "app/models",
        "app/services",
        "app/core",
        "tests"
    ]

    for folder in folders:
        (project_root / folder).mkdir(parents=True, exist_ok=True)

    llm = initialize_groq_llm()

    main_prompt = [{"role": "system", "content": "Generate the main entry point for a FastAPI application."}]
    main_code = llm.generate(main_prompt)
    if main_code.strip():
        (project_root / "app/main.py").write_text(main_code)
    else:
        print("Error: LLM did not generate code for main.py")

    routes_prompt = [{"role": "system", "content": "Generate API routes based on the extracted requirements."}]
    routes_code = llm.generate(routes_prompt)
    if routes_code.strip():
        (project_root / "app/api/routes/__init__.py").write_text(routes_code)
    else:
        print("Error: LLM did not generate code for routes.")

    models_prompt = [{"role": "system", "content": "Generate Pydantic models based on the extracted requirements."}]
    models_code = llm.generate(models_prompt)
    if models_code.strip():
        (project_root / "app/models/__init__.py").write_text(models_code)
    else:
        print("Error: LLM did not generate code for models.")

    services_prompt = [{"role": "system", "content": "Generate service functions based on the extracted requirements."}]
    services_code = llm.generate(services_prompt)
    if services_code.strip():
        (project_root / "app/services/__init__.py").write_text(services_code)
    else:
        print("Error: LLM did not generate code for services.")

    config_prompt = [{"role": "system", "content": "Generate a configuration file for the FastAPI project."}]
    config_code = llm.generate(config_prompt)
    if config_code.strip():
        (project_root / "app/core/config.py").write_text(config_code)
    else:
        print("Error: LLM did not generate code for config.")

    print("Code generation completed successfully.")

def setup_environment():
    print("Setting up virtual environment and installing dependencies...")
    project_root = Path("c:/study/autogen/generated_project")

    venv_path = project_root / "venv"
    try:
        print(f"Creating virtual environment at {venv_path}...")
        process = subprocess.Popen(
            ["python", "-m", "venv", str(venv_path)],
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE
        )
        stdout, stderr = process.communicate()
        if process.returncode != 0:
            print(f"Error creating virtual environment: {stderr.decode()}")
            raise Exception(stderr.decode())
        print(stdout.decode())
    except Exception as e:
        print(f"Error creating virtual environment: {e}")
        print(traceback.format_exc())
        raise

    pip_path = venv_path / "Scripts" / "pip"
    requirements_file = Path("c:/study/autogen/srs-analyzer/requirements.txt")

    try:
        if requirements_file.exists():
            print("Installing dependencies from requirements.txt...")
            process = subprocess.Popen(
                [str(pip_path), "install", "-r", str(requirements_file)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate()
            if process.returncode != 0:
                print(f"Error installing dependencies: {stderr.decode()}")
                raise Exception(stderr.decode())
            print(stdout.decode())
        else:
            print("requirements.txt not found. Installing default dependencies...")
            default_dependencies = ["fastapi", "uvicorn", "psycopg2", "alembic"]
            process = subprocess.Popen(
                [str(pip_path), "install"] + default_dependencies,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate()
            if process.returncode != 0:
                print(f"Error installing default dependencies: {stderr.decode()}")
                raise Exception(stderr.decode())
            print(stdout.decode())
    except Exception as e:
        print(f"Error installing dependencies: {e}")
        print(traceback.format_exc())
        raise

    print("Environment setup and dependency installation completed successfully.")

def setup_database():
    print("Setting up PostgreSQL database and migrations...")
    project_root = Path("c:/study/autogen/generated_project")

    alembic_ini = project_root / "alembic.ini"
    alembic_ini.write_text(
        """[alembic]
script_location = migrations
sqlalchemy.url = postgresql+psycopg://vectordb:vectordb@localhost:5432/vectordb

[loggers]
keys = root,sqlalchemy,alembic

[handlers]
keys = console

[formatters]
keys = generic

[logger_root]
level = WARN
handlers = console
qualname =

[logger_sqlalchemy]
level = WARN
handlers = console
qualname = sqlalchemy.engine

[logger_alembic]
level = INFO
handlers = console
qualname = alembic

[handler_console]
class = StreamHandler
args = (sys.stderr,)
level = NOTSET
formatter = generic

[formatter_generic]
format = %(asctime)s %(levelname)-5.5s [%(name)s] %(message)s
"""
    )

    migrations_path = project_root / "migrations"
    subprocess.run(["alembic", "init", str(migrations_path)], check=True, cwd=str(project_root))

    print("Database setup completed successfully.")

def validate_project():
    print("Validating the generated project...")
    project_root = Path("c:/study/autogen/generated_project")

    main_file = project_root / "app/main.py"
    if main_file.exists():
        print("Starting the generated project on localhost:8001...")
        try:
            process = subprocess.Popen(
                ["uvicorn", "app.main:app", "--host", "127.0.0.1", "--port", "8001", "--reload"],
                cwd=str(project_root),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE
            )
            stdout, stderr = process.communicate()
            if process.returncode != 0:
                print(f"Error starting the project: {stderr.decode()}")
                raise Exception(stderr.decode())
            print(stdout.decode())
            print("Project is running at http://127.0.0.1:8001")
        except Exception as e:
            print(f"Error starting the project: {e}")
            print(traceback.format_exc())
            raise
    else:
        print("Main file not found. Validation failed.")
        raise FileNotFoundError("Main file not found in the generated project.")

    print("Project validation completed successfully.")
